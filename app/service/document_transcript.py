"""把导入的文本文档转成飞书妙记同款的转写格式。

下游的一切（纪要提示词、时间锚点、划词提问、参会人解析）都按那个格式读，所以
不管进来的是 txt 还是 pdf，出去都得长成 `说话人1 HH:MM:SS.mmm` 加正文的样子。

纯文本没有真实时间轴，只有 srt 例外。没有时间轴时按段落顺序编一条递增的假时间轴：
锚点跳转会落到大致对应的段落上，比全部堆在 00:00:00 有用得多。
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path

from app.service.asr_transcript import (
    TranscriptLine,
    format_clock,
    format_duration,
    local_label,
)

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = {".txt", ".md", ".srt", ".docx", ".pdf"}

# 没有真实时间轴时，假定每段读这么久，且每段至少占这么久
SECONDS_PER_CHAR = 0.25
MIN_SEGMENT_SECONDS = 3.0

SRT_TIME_RE = re.compile(
    r"(\d{1,2}):(\d{2}):(\d{2})[,.](\d{1,3})\s*-->\s*"
    r"(\d{1,2}):(\d{2}):(\d{2})[,.](\d{1,3})"
)


class DocumentParseError(RuntimeError):
    """文档读不出可用文本。"""


@dataclass
class ParsedDocument:
    transcript: str
    # srt 能给出真实时长；其它格式是按字数估的
    duration_ms: int
    exact_duration: bool
    segments: int


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".srt":
        return _from_srt(_read_text(path))
    if suffix == ".docx":
        return _from_paragraphs(_read_docx(path))
    if suffix == ".pdf":
        return _from_paragraphs(_read_pdf(path))
    if suffix in {".txt", ".md"}:
        return _from_paragraphs(_split_paragraphs(_read_text(path)))
    raise DocumentParseError(f"不支持的文本格式：{suffix or '（无扩展名）'}")


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "utf-16"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    # 全试完还不行就强解一遍，个别乱码字符不该让整份文档进不来
    return raw.decode("utf-8", errors="replace")


def _read_docx(path: Path) -> list[str]:
    try:
        import docx
    except ImportError as exc:
        raise DocumentParseError(
            "未安装 python-docx，无法读取 .docx；请先 pip install python-docx"
        ) from exc
    document = docx.Document(str(path))
    blocks = [p.text.strip() for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append("　".join(cells))
    return [b for b in blocks if b]


def _read_pdf(path: Path) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError(
            "未安装 pypdf，无法读取 .pdf；请先 pip install pypdf"
        ) from exc
    reader = PdfReader(str(path))
    blocks: list[str] = []
    for page in reader.pages:
        blocks.extend(_split_paragraphs(page.extract_text() or ""))
    return blocks


def _split_paragraphs(text: str) -> list[str]:
    """按空行分段；整篇没有空行时退化成按行分。"""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    blocks = [b.strip() for b in re.split(r"\n\s*\n", normalized)]
    blocks = [b for b in blocks if b]
    if len(blocks) <= 1:
        blocks = [line.strip() for line in normalized.split("\n") if line.strip()]
    return blocks


def _estimate_seconds(text: str) -> float:
    return max(MIN_SEGMENT_SECONDS, len(text) * SECONDS_PER_CHAR)


def _render(lines: list[TranscriptLine], *, duration_ms: int) -> str:
    """按飞书原文的形状渲染。

    不走 `build_transcript`：它会把同一说话人的相邻段落并成一大块，而导入的文档
    全篇只有「说话人1」，并完之后段落边界连同时间轴一起没了，锚点就再没处落。
    """
    body = [format_duration(duration_ms), ""]
    for line in lines:
        # 时间戳后面那个空格是飞书原文的一部分，段头正则依赖这行的整体形状
        body.append(f"{line.speaker} {format_clock(line.start_ms)} ")
        body.append(line.text)
        body.append("")
    return "\n".join(body).rstrip() + "\n"


def _from_paragraphs(blocks: list[str]) -> ParsedDocument:
    """按段落编一条递增的假时间轴。"""
    blocks = [b for b in (x.strip() for x in blocks) if b]
    if not blocks:
        raise DocumentParseError("文档里没有可用的文字内容")

    speaker = local_label(1)
    lines: list[TranscriptLine] = []
    cursor = 0.0
    for block in blocks:
        length = _estimate_seconds(block)
        lines.append(
            TranscriptLine(
                speaker=speaker,
                start_ms=int(cursor * 1000),
                end_ms=int((cursor + length) * 1000),
                text=block,
            )
        )
        cursor += length
    duration_ms = int(cursor * 1000)
    return ParsedDocument(
        transcript=_render(lines, duration_ms=duration_ms),
        duration_ms=duration_ms,
        exact_duration=False,
        segments=len(lines),
    )


def _from_srt(text: str) -> ParsedDocument:
    """srt 自带时间轴，直接用；说话人一律记成「说话人1」。"""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines: list[TranscriptLine] = []
    speaker = local_label(1)
    for chunk in re.split(r"\n\s*\n", normalized):
        match = SRT_TIME_RE.search(chunk)
        if match is None:
            continue
        body_lines = [
            line.strip()
            for line in chunk.split("\n")
            if line.strip() and not SRT_TIME_RE.search(line) and not line.strip().isdigit()
        ]
        body = " ".join(body_lines).strip()
        if not body:
            continue
        lines.append(
            TranscriptLine(
                speaker=speaker,
                start_ms=_srt_ms(match, 0),
                end_ms=_srt_ms(match, 4),
                text=body,
            )
        )
    if not lines:
        # 时间轴解析不出来时按普通文本处理，总比整份丢掉好
        logger.info("srt 里没有解析出时间轴，退化成按段落导入")
        return _from_paragraphs(_split_paragraphs(text))
    duration_ms = max(line.end_ms for line in lines)
    return ParsedDocument(
        transcript=_render(lines, duration_ms=duration_ms),
        duration_ms=duration_ms,
        exact_duration=True,
        segments=len(lines),
    )


def _srt_ms(match: re.Match[str], offset: int) -> int:
    hours = int(match.group(offset + 1))
    minutes = int(match.group(offset + 2))
    seconds = int(match.group(offset + 3))
    millis_text = match.group(offset + 4)
    millis = int(millis_text.ljust(3, "0"))
    return ((hours * 60 + minutes) * 60 + seconds) * 1000 + millis
