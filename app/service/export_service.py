"""会议纪要 / 转写导出。"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path

from app.core import runtime_config
from app.service.meeting_storage_service import MeetingStorageService

logger = logging.getLogger(__name__)

IMAGE_LINE_RE = re.compile(r"^!\[[^\]]*\]\(\s*([^)\s]+)\s*\)$")
CAPTION_LINE_RE = re.compile(r"^图：")


class ExportService:
    def __init__(self) -> None:
        self._storage = MeetingStorageService()

    def export_summary(
        self, minute_token: str, fmt: str, *, owner_user_id: int
    ) -> tuple[bytes, str, str]:
        fmt = fmt.lower()
        detail = self._storage.read_summary(
            minute_token, owner_user_id=owner_user_id
        )
        if detail is None:
            raise LookupError("该会议尚未生成纪要")
        content = detail["content"] or ""
        title = (
            self._storage.read_meta(minute_token, owner_user_id=owner_user_id) or {}
        ).get("title") or minute_token
        safe = self._safe_filename(title)
        watermark = self._watermark_text()

        if fmt == "md":
            text = self.strip_images(content)
            return text.encode("utf-8"), f"{safe}-summary.md", "text/markdown; charset=utf-8"
        if fmt == "docx":
            data = self._summary_to_docx(
                minute_token,
                content,
                title,
                owner_user_id=owner_user_id,
                watermark=watermark,
            )
            return data, f"{safe}-summary.docx", (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        if fmt == "pdf":
            data = self._summary_to_pdf(
                minute_token,
                content,
                title,
                owner_user_id=owner_user_id,
                watermark=watermark,
            )
            return data, f"{safe}-summary.pdf", "application/pdf"
        raise ValueError("纪要导出格式仅支持 pdf / docx / md")

    def export_transcript(
        self, minute_token: str, fmt: str, *, owner_user_id: int
    ) -> tuple[bytes, str, str]:
        fmt = fmt.lower()
        text = self._storage.read_transcript(
            minute_token, owner_user_id=owner_user_id
        )
        if text is None:
            raise LookupError("本地没有该会议的转写文本")
        title = (
            self._storage.read_meta(minute_token, owner_user_id=owner_user_id) or {}
        ).get("title") or minute_token
        safe = self._safe_filename(title)

        if fmt == "txt":
            return text.encode("utf-8"), f"{safe}-transcript.txt", "text/plain; charset=utf-8"
        if fmt == "md":
            md = f"# {title}\n\n```\n{text.rstrip()}\n```\n"
            return md.encode("utf-8"), f"{safe}-transcript.md", "text/markdown; charset=utf-8"
        raise ValueError("转写导出格式仅支持 md / txt")

    @staticmethod
    def _watermark_text() -> str:
        return (runtime_config.get_str("EXPORT_WATERMARK_TEXT", "") or "").strip()

    @staticmethod
    def strip_images(markdown: str) -> str:
        out: list[str] = []
        skip_caption = False
        for raw in markdown.splitlines():
            line = raw.rstrip()
            if IMAGE_LINE_RE.match(line.strip()):
                skip_caption = True
                continue
            if skip_caption and CAPTION_LINE_RE.match(line.strip()):
                skip_caption = False
                continue
            skip_caption = False
            out.append(raw)
        return "\n".join(out).strip() + "\n"

    def _resolve_image_path(
        self, minute_token: str, ref: str, *, owner_user_id: int
    ) -> Path | None:
        name = ref.replace("\\", "/").lstrip("./")
        if name.startswith("assets/"):
            name = name[len("assets/") :]
        if not name or "/" in name or "\\" in name:
            return None
        local = self._storage.resolve_asset_path(
            minute_token, name, owner_user_id=owner_user_id
        )
        if local is not None:
            return local
        from app.core.async_bridge import run_async
        from app.service.r2_media_service import r2_media_service

        try:
            return run_async(
                r2_media_service.materialize_asset(
                    minute_token, name, owner_user_id=owner_user_id
                )
            )
        except Exception as exc:  # noqa: BLE001
            logger.error(
                "导出时从 R2 物化配图失败 token=%s file=%s，该图将跳过。err=%s",
                minute_token,
                name,
                exc,
            )
            return None

    def _summary_to_docx(
        self,
        minute_token: str,
        markdown: str,
        title: str,
        *,
        owner_user_id: int,
        watermark: str = "",
    ) -> bytes:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()
        if watermark:
            for section in doc.sections:
                header = section.header
                header.is_linked_to_previous = False
                paragraph = (
                    header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                )
                paragraph.text = ""
                run = paragraph.add_run(watermark)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xB0, 0xB0, 0xB0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading(title, level=0)
        for raw in markdown.splitlines():
            line = raw.rstrip()
            stripped = line.strip()
            if not stripped:
                continue
            image = IMAGE_LINE_RE.match(stripped)
            if image:
                path = self._resolve_image_path(
                    minute_token, image.group(1), owner_user_id=owner_user_id
                )
                if path and path.is_file():
                    try:
                        doc.add_picture(str(path), width=Inches(5.5))
                    except Exception:
                        logger.warning(
                            "纪要导出 docx 嵌入图片失败 token=%s path=%s，怀疑文件损坏，将跳过该图",
                            minute_token,
                            path,
                        )
                continue
            if CAPTION_LINE_RE.match(stripped):
                doc.add_paragraph(stripped)
                continue
            heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if heading:
                level = min(len(heading.group(1)), 4)
                doc.add_heading(heading.group(2), level=level)
                continue
            if stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
                continue
            ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
            if ordered:
                doc.add_paragraph(ordered.group(1), style="List Number")
                continue
            doc.add_paragraph(stripped)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _summary_to_pdf(
        self,
        minute_token: str,
        markdown: str,
        title: str,
        *,
        owner_user_id: int,
        watermark: str = "",
    ) -> bytes:
        import markdown as md_lib
        from xhtml2pdf import pisa

        html_body = md_lib.markdown(
            markdown,
            extensions=["extra", "sane_lists"],
            output_format="html5",
        )

        def replace_img(match: re.Match[str]) -> str:
            src = match.group(1)
            path = self._resolve_image_path(
                minute_token, src, owner_user_id=owner_user_id
            )
            if path and path.is_file():
                return f'<img src="{path.as_uri()}" style="max-width:100%;" />'
            return ""

        html_body = re.sub(
            r'<img[^>]+src=["\']([^"\']+)["\'][^>]*>',
            replace_img,
            html_body,
            flags=re.IGNORECASE,
        )
        watermark_html = ""
        if watermark:
            wm = _escape_html(watermark)
            watermark_html = f"""
<div id="wm1" style="position: fixed; top: 28%; left: 8%; opacity: 0.14; font-size: 26pt;
  color: #888; transform: rotate(-28deg); z-index: 0;">{wm}</div>
<div id="wm2" style="position: fixed; top: 58%; left: 22%; opacity: 0.12; font-size: 22pt;
  color: #888; transform: rotate(-28deg); z-index: 0;">{wm}</div>
"""
        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"/>
<style>
body {{ font-family: "Microsoft YaHei", "PingFang SC", sans-serif; font-size: 12pt; line-height: 1.5; }}
h1,h2,h3 {{ margin-top: 1.2em; }}
img {{ max-width: 100%; }}
code {{ font-family: Consolas, monospace; }}
.content {{ position: relative; z-index: 1; }}
</style></head>
<body>
{watermark_html}
<div class="content"><h1>{_escape_html(title)}</h1>{html_body}</div>
</body></html>"""

        buffer = io.BytesIO()
        result = pisa.CreatePDF(html, dest=buffer, encoding="utf-8")
        if result.err:
            raise RuntimeError("PDF 生成失败，怀疑 xhtml2pdf 无法解析当前纪要 HTML")
        return buffer.getvalue()

    @staticmethod
    def _safe_filename(name: str) -> str:
        # HTTP Content-Disposition 头只能用 latin-1；中文标题改成可读 ASCII
        cleaned = re.sub(r"[^\w.\-]+", "_", name, flags=re.UNICODE).strip("._") or "export"
        ascii_name = cleaned.encode("ascii", "ignore").decode("ascii").strip("._")
        return (ascii_name or "export")[:80]


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


export_service = ExportService()
